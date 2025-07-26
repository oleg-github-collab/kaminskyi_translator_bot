import os
import logging
from docx import Document
import pdfplumber
import tempfile

logger = logging.getLogger(__name__)

def count_chars_in_file(file_path: str) -> int:
    """МАКСИМАЛЬНО НАДІЙНИЙ підрахунок символів у файлі"""
    try:
        logger.info(f"Starting character count for file: {file_path}")
        
        # Перевірка існування файлу
        if not os.path.exists(file_path):
            logger.error(f"File does not exist: {file_path}")
            return 0
            
        # Перевірка чи файл не порожній
        if os.path.getsize(file_path) == 0:
            logger.warning(f"File is empty: {file_path}")
            return 0
            
        # Отримання розширення файлу
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        logger.info(f"File extension: {ext} for file: {file_path}")
        
        char_count = 0
        
        if ext == '.txt':
            char_count = _count_txt_file(file_path)
        elif ext == '.docx':
            char_count = _count_docx_file(file_path)
        elif ext == '.pdf':
            char_count = _count_pdf_file(file_path)
        else:
            logger.warning(f"Unsupported file type: {ext}")
            return 0
            
        logger.info(f"Final character count for {file_path}: {char_count}")
        return char_count
        
    except Exception as e:
        logger.error(f"Critical error in count_chars_in_file for {file_path}: {str(e)}", exc_info=True)
        return 0

def _count_txt_file(file_path: str) -> int:
    """Підрахунок символів у TXT файлі"""
    try:
        encodings = ['utf-8', 'windows-1251', 'cp1251', 'latin1', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                    char_count = len(content)
                    if char_count > 0:
                        logger.info(f"Successfully read TXT file {file_path} with {encoding}: {char_count} chars")
                        return char_count
            except Exception as e:
                logger.debug(f"Failed to read with {encoding}: {str(e)}")
                continue
                
        logger.warning(f"Could not read meaningful content from TXT file: {file_path}")
        return 0
        
    except Exception as e:
        logger.error(f"Error counting TXT file {file_path}: {str(e)}")
        return 0

def _count_docx_file(file_path: str) -> int:
    """Підрахунок символів у DOCX файлі"""
    try:
        doc = Document(file_path)
        total_chars = 0
        for paragraph in doc.paragraphs:
            total_chars += len(paragraph.text)
        logger.info(f"Successfully counted DOCX file {file_path}: {total_chars} chars")
        return total_chars
    except Exception as e:
        logger.error(f"Error counting DOCX file {file_path}: {str(e)}")
        return 0

def _count_pdf_file(file_path: str) -> int:
    """Підрахунок символів у PDF файлі"""
    try:
        total_chars = 0
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    if text:
                        total_chars += len(text)
                        logger.debug(f"Page {page_num} extracted {len(text)} chars")
                except Exception as page_error:
                    logger.warning(f"Error extracting text from page {page_num} in {file_path}: {str(page_error)}")
                    continue
        logger.info(f"Successfully counted PDF file {file_path}: {total_chars} chars")
        return total_chars
    except Exception as e:
        logger.error(f"Error counting PDF file {file_path}: {str(e)}")
        return 0

def read_file_content(file_path: str) -> str:
    """Надійне читання вмісту файлу"""
    try:
        if not os.path.exists(file_path):
            logger.warning(f"File not found for reading: {file_path}")
            return ""
            
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        content = ""
        
        if ext == '.txt':
            # Пробуємо різні кодування
            encodings = ['utf-8', 'windows-1251', 'cp1251', 'latin1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        content = f.read()
                        if content.strip():  # Якщо файл не порожній
                            logger.info(f"Successfully read TXT file with {encoding}")
                            break
                except Exception:
                    continue
            if not content.strip():
                logger.warning(f"Could not read meaningful content from TXT file: {file_path}")
                
        elif ext == '.docx':
            try:
                doc = Document(file_path)
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text)
                content = '\n'.join(paragraphs)
                logger.info(f"Successfully read DOCX file: {file_path}")
            except Exception as e:
                logger.error(f"Error reading DOCX file {file_path}: {str(e)}")
                
        elif ext == '.pdf':
            try:
                text = ''
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + '\n'
                        except Exception as page_error:
                            logger.warning(f"Error extracting page from PDF {file_path}: {str(page_error)}")
                            continue
                content = text.strip()
                logger.info(f"Successfully read PDF file: {file_path}")
            except Exception as e:
                logger.error(f"Error reading PDF file {file_path}: {str(e)}")
                
        return content
        
    except Exception as e:
        logger.error(f"Unexpected error reading file {file_path}: {str(e)}")
        return ""

def write_translated_file(file_path: str, translated_text: str, original_ext: str) -> str:
    """Надійне створення перекладеного файлу"""
    try:
        if not translated_text or not translated_text.strip():
            logger.warning(f"Translated text is empty for file: {file_path}")
            raise Exception("Перекладений текст порожній")
            
        base_name = os.path.splitext(file_path)[0]
        new_file_path = f"{base_name}_translated{original_ext}"
        
        logger.info(f"Writing translated file: {new_file_path}")
        
        if original_ext == '.txt':
            with open(new_file_path, 'w', encoding='utf-8') as f:
                f.write(translated_text)
            logger.info(f"Created translated TXT file: {new_file_path}")
                
        elif original_ext == '.docx':
            doc = Document()
            lines = translated_text.split('\n')
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line)
            doc.save(new_file_path)
            logger.info(f"Created translated DOCX file: {new_file_path}")
                
        elif original_ext == '.pdf':
            # Для PDF зберігаємо як TXT
            txt_path = new_file_path.replace('.pdf', '.txt')
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(translated_text)
            logger.info(f"Created translated TXT file from PDF: {txt_path}")
            return txt_path
        else:
            # За замовчуванням зберігаємо як TXT
            txt_path = new_file_path + '.txt'
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(translated_text)
            logger.info(f"Created default TXT file: {txt_path}")
            return txt_path
            
        return new_file_path
        
    except Exception as e:
        logger.error(f"Error writing translated file {file_path}: {str(e)}", exc_info=True)
        raise Exception(f"Помилка створення файлу: {str(e)}")

def cleanup_temp_file(file_path: str):
    """Надійне очищення тимчасових файлів"""
    try:
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"Cleaned up temporary file: {file_path}")
            return True
        return False
    except Exception as e:
        logger.error(f"Error cleaning up file {file_path}: {str(e)}")
        return False